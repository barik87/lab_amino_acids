import glob

from docx import Document
from openpyxl import Workbook
from openpyxl import load_workbook
from openpyxl.cell import Cell
from openpyxl.styles import PatternFill

import dictionaries

import argparse

parser = argparse.ArgumentParser()
parser.add_argument('--mult', dest='is_mult', action='store_const', const=True, default=False,
                    help='whether to multiply some amino acids')
parser.add_argument('--all-mult', dest='is_all_mult', action='store_const', const=True, default=False,
                    help='whether to multiply all amino acids')
args = parser.parse_args()

DATA_FILE_NAME = 'data.xlsx'
TEMPLATE_PLAZMA_FILE_NAME = 'template_plazma.docx'
TEMPLATE_URINE_FILE_NAME = 'template_urine.docx'

MULTS = dictionaries.MULTIPLIERS
ALL_MULTS = dictionaries.CONST_MULTIPLIERS
PARAMETERS = dictionaries.get_parameters_from_file(DATA_FILE_NAME)

DATA = dict()
print('Reading data from "%s"...' % DATA_FILE_NAME)
wb = load_workbook(DATA_FILE_NAME)
ws = wb.worksheets[0]
headers = list(map(lambda row_cell: str(row_cell.value), list(ws.rows)[0]))
for row_idx in range(1, len(list(ws.rows))):
    row = list(ws.rows)[row_idx]
    row_id = None
    for col_idx in range(len(row)):
        cell = row[col_idx]
        if col_idx == 0:
            row_id = str(cell.value)
            DATA[row_id] = dict()
        else:
            DATA[row_id][headers[col_idx]] = cell.value
print('Done.')

green_fill = PatternFill("solid", fgColor='92D050')
yellow_fill = PatternFill("solid", fgColor='FFC000')
red_fill = PatternFill("solid", fgColor='FF0000')

# create total_results.xlsx
acid_list = dictionaries.get_acids_list(DATA_FILE_NAME)
total_wb = Workbook()
total_ws = total_wb.active
acids_row = [''] + acid_list
total_ws.append(acids_row)
corrected_acids = set()

for file_name in glob.glob('*.xlsx'):
    if 'results' not in file_name and file_name != DATA_FILE_NAME:
        print('\nWorking with file "%s"...' % file_name)

        file_id = file_name.split('.')[0]
        data = DATA[file_id]

        norm_age, acid_norms = dictionaries.get_acids(DATA_FILE_NAME, data['Type'], data['SampleTakeDate'],
                                                      data['BirthDate'])

        wb = load_workbook(file_name)
        ws = wb.worksheets[0]

        start_row_idx = -1
        start_col_idx = -1
        end_row_idx = -1
        end_col_idx = -1
        area_col_idx = -1
        for row_idx in range(len(list(ws.rows))):
            row = list(ws.rows)[row_idx]
            if start_row_idx == -1:
                for col_idx in range(len(row)):
                    cell = row[col_idx]
                    if str(cell.value) == 'Peak Name':
                        start_row_idx = row_idx
                        start_col_idx = col_idx
                    if str(cell.value).strip() == 'Area':
                        area_col_idx = col_idx
                    if str(cell.value).strip() == 'Amount':
                        end_col_idx = col_idx
                        break
            elif row_idx == len(list(ws.rows)) - 1:
                end_row_idx = row_idx
            elif row_idx > start_row_idx + 2:
                for col_idx in range(len(row)):
                    if row[col_idx].value is not None:
                        break
                else:
                    end_row_idx = row_idx - 1
            if start_row_idx > -1 and start_col_idx > -1 and end_row_idx > -1 and end_col_idx > -1:
                break
        else:
            print('ERROR: Cell with text "Peak Name" not found in the document.')

        amounts = dict()
        total_row = [data['Name']] + [''] * len(acid_list)
        new_wb = Workbook()
        new_ws = new_wb.active
        for row_idx in range(start_row_idx, end_row_idx + 1):
            row = list(ws.rows)[row_idx]
            acid_name = row[start_col_idx].value
            if acid_name is not None and acid_name in acid_list:
                new_row = []
                for col_idx in range(start_col_idx, end_col_idx + 1):
                    cell = row[col_idx]
                    value = '' if cell.value is None else cell.value
                    if col_idx == end_col_idx:
                        if args.is_all_mult and row_idx > start_row_idx + 2:
                            try:
                                area_value = float(row[area_col_idx].value)
                            except ValueError:
                                area_value = 0
                            value = area_value / ALL_MULTS[str(row[start_col_idx].value)]
                        elif value == 'n.a.':
                            if row[area_col_idx].value != 'n.a.':
                                if args.is_mult:
                                    value = row[area_col_idx].value * MULTS[data['Type']][str(row[start_col_idx].value)]
                                else:
                                    value = row[area_col_idx].value
                            else:
                                value = 0
                        if str(PARAMETERS['CorrectResults']).lower() == 'yes' and row_idx > start_row_idx + 2:
                            divider = float(PARAMETERS[str(row[start_col_idx].value) + 'Divider'])
                            if divider != 1.0:
                                value = value / divider
                                corrected_acids.add(str(row[start_col_idx].value))
                        if type(value) in [int, float]:
                            value = value * PARAMETERS['Multiplier']
                        if data['Type'] == 'Urine' and value != 'n.a.' and row_idx > start_row_idx + 2:
                            try:
                                value = float(value) / float(data['Creatinine'])
                            except TypeError:
                                print('ERROR: Creatinine value for ID %s is not number.' % file_id)
                        amounts[acid_name] = value
                        if acid_name in acid_norms:
                            new_cell = Cell(new_ws)
                            new_cell.value = value
                            total_cell = Cell(total_ws)
                            total_cell.value = value

                            lbound, rbound = acid_norms[acid_name]
                            if float(value) < float(lbound):
                                new_cell.fill = yellow_fill
                                total_cell.fill = yellow_fill
                            elif float(value) <= float(rbound):
                                new_cell.fill = green_fill
                                total_cell.fill = green_fill
                            else:
                                new_cell.fill = red_fill
                                total_cell.fill = red_fill

                            new_row.append(new_cell)
                            new_row.append('%s--%s' % (lbound, rbound))
                            total_row[1 + acid_list.index(acid_name)] = total_cell
                        else:
                            new_row.append(value)
                    else:
                        new_row.append(value)
                new_ws.append(new_row)
        new_file_name = file_id + '_results.xlsx'
        new_wb.save(filename=new_file_name)
        print('File %s generated.' % new_file_name)

        total_ws.append(total_row)

        temp_doc_name = TEMPLATE_PLAZMA_FILE_NAME if data['Type'] == 'Plazma' else TEMPLATE_URINE_FILE_NAME
        document = Document(temp_doc_name)
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                if 'FULLNAME' in run.text:
                    run.text = str(run.text).replace('FULLNAME', data['Name'])
                if 'DDMMYYYY' in run.text:
                    run.text = str(run.text).replace('DDMMYYYY', data['BirthDate'].strftime('%d.%m.%Y'))
                if 'SAMPLEDAY' in run.text:
                    run.text = str(run.text).replace('SAMPLEDAY', data['SampleTakeDate'].strftime('%d'))
                if 'SAMPLEYEAR' in run.text:
                    run.text = str(run.text).replace('SAMPLEYEAR', data['SampleTakeDate'].strftime('%Y'))
                if 'SAMPLEMONTH' in run.text:
                    run.text = str(run.text).replace('SAMPLEMONTH',
                                                     dictionaries.UKR_MONTHS[data['SampleTakeDate'].month])
                if 'ANDAY' in run.text:
                    run.text = str(run.text).replace('ANDAY', data['AnalysisDate'].strftime('%d'))
                if 'ANYEAR' in run.text:
                    run.text = str(run.text).replace('ANYEAR', data['AnalysisDate'].strftime('%Y'))
                if 'ANMONTH' in run.text:
                    run.text = str(run.text).replace('ANMONTH',
                                                     dictionaries.UKR_MONTHS[data['AnalysisDate'].month])

        table = document.tables[0]
        for col_idx in range(3, 5):
            for p in table.cell(0, col_idx).paragraphs:
                for r in p.runs:
                    if r.text == 'BOUND':
                        r.text = norm_age

        for row_idx in range(1, len(table.rows)):
            acid_name = acid_list[row_idx - 1]
            lbound, rbound = acid_norms[acid_name]

            amount = float(amounts[acid_name])
            suffix = ''
            # if amount < float(lbound):
            # suffix = '↓'
            # elif amount > float(rbound):
            #     suffix = '↑'

            str_amount = ('0' if amounts[acid_name] == 0 else '%.2f' % amount) + suffix
            str_amount = str_amount.replace('.', ',')
            table.cell(row_idx, 1).paragraphs[0].runs[0].text = str_amount
            table.cell(row_idx, 3).paragraphs[0].runs[0].text = str(lbound)
            table.cell(row_idx, 4).paragraphs[0].runs[0].text = str(rbound)

        new_doc_file_name = '_'.join(list(data['Name'].split(' ')) + [data['Type']]) + '.docx'
        document.save(new_doc_file_name)
        print('File docx generated.')

total_file_name = 'total_results.xlsx'
total_wb.save(filename=total_file_name)
print('\nFile %s generated.' % total_file_name)

if str(PARAMETERS['CorrectResults']).lower() == 'yes':
    print('\n======================================= ATTENTION =======================================')
    print('!!! Results were corrected: %s !!!' % ', '.join(corrected_acids))
    print('======================================= ATTENTION =======================================')
